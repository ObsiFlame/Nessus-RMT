from docx import Document
from docx.oxml import OxmlElement, ns
from datetime import datetime
import os


# -----------------------------------------------------------
# Severity mapping
# -----------------------------------------------------------
SEVERITY_MAP = {
    4: "Critical",
    3: "High",
    2: "Medium",
    1: "Low",
    0: "Info"
}

SEVERITY_COLOR = {
    "Critical": "FF0000",  # Red
    "High": "FFA500",      # Orange
    "Medium": "FFFF00",    # Yellow
    "Low": "00FF00",       # Green
    "Info": "D3D3D3"       # Gray
}


# -----------------------------------------------------------
# Helpers
# -----------------------------------------------------------
def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_row(table, key, value, bg_color=None):
    row = table.add_row().cells

    # Header cell (bold)
    key_para = row[0].paragraphs[0]
    key_run = key_para.add_run(key)
    key_run.bold = True

    # Value cell
    val_para = row[1].paragraphs[0]
    val_run = val_para.add_run(str(value if value else "N/A"))

    # Apply background color to BOTH cells if needed
    if bg_color:
        set_cell_background(row[0], bg_color)
        set_cell_background(row[1], bg_color)


# -----------------------------------------------------------
# DOCX Generator
# -----------------------------------------------------------
def generate_docx(scan_id, findings):
    os.makedirs("reports", exist_ok=True)

    filename = f"reports/scan_{scan_id}_report.docx"
    doc = Document()

    # Title
    doc.add_heading(
        f"Vulnerability Assessment Report – Scan {scan_id}",
        level=1
    )

    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    for idx, finding in enumerate(findings, start=1):


        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Resolve severity
        sev_numeric = finding.get("severity")
        sev_text = SEVERITY_MAP.get(sev_numeric, "Unknown")
        sev_color = SEVERITY_COLOR.get(sev_text)

        # Rows (Plugin ID REMOVED)
        add_row(table, "Sl No", idx)
        add_row(table, "Affected Host", finding.get("hosts"))
        add_row(table, "Vulnerability Title", finding.get("plugin_name"))

        # Severity row (FULL COLOR)
        add_row(
            table,
            "Severity",
            sev_text,
            bg_color=sev_color
        )

        add_row(table, "CVE / CWE", finding.get("cves"))
        add_row(table, "Description", finding.get("description"))
        add_row(table, "Recommendation", finding.get("solution"))
        add_row(table, "Proof of Concept", finding.get("plugin_output"))
        add_row(table, "Business Impact", finding.get("synopsis"))
        add_row(table, "References", finding.get("references"))

    doc.save(filename)
    return filename
